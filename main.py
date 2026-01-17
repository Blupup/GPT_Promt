import os
import re
from pathlib import Path
from typing import Dict, Optional
from openai import OpenAI
from datetime import datetime
from docx import Document

SYSTEM_INSTRUCTION = """You are a film director, anthropologist, and visual historian creating cinematic video prompts for Google Veo 3 (fast mode).
Your task is to generate 1 prompt in English from the provided paragraph of a prehistoric narrative script.

Guidelines:
- Create vivid, cinematic descriptions suitable for video generation
- Focus on visual elements, atmosphere, and camera work
- Keep prompts concise but descriptive (2-4 sentences)
- Include details about lighting, composition, and mood
- Respond ONLY with the video prompt text, no additional commentary"""


def parse_document(file_path: str | Path) -> Dict[int, str]:
    file_path = Path(file_path)

    if file_path.suffix.lower() == '.docx':
        doc = Document(file_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        content = file_path.read_text(encoding='utf-8')

    print(f"\n[DEBUG] Первые 500 символов документа:")
    print("=" * 80)
    print(content[:500])
    print("=" * 80)
    print(f"\n[DEBUG] Всего символов в документе: {len(content)}\n")

    pattern = r'(\d+)\.\s+(.+?)(?=\n\d+\.\s+|\Z)'
    matches = re.findall(pattern, content, re.DOTALL)

    paragraphs = {}

    if matches:
        for num, text in matches:
            clean_text = ' '.join(text.split())
            paragraphs[int(num)] = clean_text
    else:
        print("[INFO] Нумерация не найдена. Разбиваем документ по абзацам...")
        if file_path.suffix.lower() == '.docx':
            doc = Document(file_path)
            para_num = 1
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > 20:
                    paragraphs[para_num] = text
                    para_num += 1
        else:
            parts = content.split('\n\n')
            para_num = 1
            for part in parts:
                text = ' '.join(part.split()).strip()
                if text and len(text) > 20:
                    paragraphs[para_num] = text
                    para_num += 1

    return paragraphs

def generate_video_prompts(
        paragraphs: Dict[int, str],
        api_key: str,
        site_url: str = "",
        site_name: str = "Video Prompt Generator",
        max_paragraphs: Optional[int] = None
) -> Dict[int, str]:
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
    )

    results = {}

    items = list(paragraphs.items())
    if max_paragraphs:
        items = items[:max_paragraphs]

    for num, paragraph_text in items:
        print(f"\nОбработка абзаца {num}...")
        print(f"Текст: {paragraph_text[:100]}...")

        try:
            extra_headers = {}
            if site_url:
                extra_headers["HTTP-Referer"] = site_url
            if site_name:
                extra_headers["X-Title"] = site_name

            completion = client.chat.completions.create(
                extra_headers=extra_headers,
                extra_body={},
                model="openai/gpt-4o",
                messages=[
                    {"role": "system", "content": SYSTEM_INSTRUCTION},
                    {"role": "user", "content": paragraph_text}
                ],
                temperature=0.7,
                max_tokens=200
            )

            video_prompt = completion.choices[0].message.content.strip()
            results[num] = video_prompt
            print(f"✓ Промпт сгенерирован: {video_prompt[:80]}...")

        except Exception as e:
            print(f"✗ Ошибка при обработке абзаца {num}: {e!r}")
            results[num] = f"ERROR: {e!r}"

    return results


def save_results_txt(results: Dict[int, str], output_file: str | Path) -> None:
    output_file = Path(output_file)

    with output_file.open('w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("VIDEO PROMPTS GENERATED\n")
        f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 80 + "\n\n")

        for num in sorted(results.keys()):
            f.write(f"PARAGRAPH {num}\n")
            f.write("-" * 80 + "\n")
            f.write(f"{results[num]}\n\n")


def save_results_csv(results: Dict[int, str], output_file: str | Path) -> None:
    import csv

    output_file = Path(output_file)

    with output_file.open('w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['Paragraph Number', 'Video Prompt'])

        for num in sorted(results.keys()):
            writer.writerow([num, results[num]])


def main() -> None:

    INPUT_FILE = "200. How People Kept Warm in Castles Without Fireplaces ENG.docx"
    OUTPUT_TXT = "video_prompts.txt"
    OUTPUT_CSV = "video_prompts.csv"
    MAX_PARAGRAPHS = 2 #поставил для теста, если нужно весь файл прогнать напишите None

    OPENROUTER_API_KEY = "" # сюда вставите свой API
    SITE_URL = os.getenv("SITE_URL", "")
    SITE_NAME = os.getenv("SITE_NAME", "Video Prompt Generator")

    if not OPENROUTER_API_KEY:
        print("=" * 80)
        print("ОШИБКА: Не найден API ключ OpenRouter!")
        print("=" * 80)
        print("\nУстановите переменную окружения OPENROUTER_API_KEY")
        print("=" * 80)
        return

    print("VIDEO PROMPT GENERATOR")
    print(f"Python version: {os.sys.version}")
    print(f"Input file: {INPUT_FILE}")
    print(f"Processing: {MAX_PARAGRAPHS or 'ALL'} paragraphs")
    print("=" * 80)

    print(f"\n[1/3] Чтение файла...")
    try:
        paragraphs = parse_document(INPUT_FILE)
        print(f"      Найдено абзацев: {len(paragraphs)}")
    except FileNotFoundError:
        print(f"      ОШИБКА: Файл не найден")
        return
    except Exception as e:
        print(f"      ОШИБКА при чтении: {e!r}")
        return

    print(f"\n[2/3] Генерация видео-промптов...")
    results = generate_video_prompts(
        paragraphs,
        OPENROUTER_API_KEY,
        site_url=SITE_URL,
        site_name=SITE_NAME,
        max_paragraphs=MAX_PARAGRAPHS
    )
    print(f"\n      Обработано абзацев: {len(results)}")

    print(f"\n[3/3] Сохранение результатов...")
    save_results_txt(results, OUTPUT_TXT)
    print(f"      Текстовый файл: {OUTPUT_TXT}")

    save_results_csv(results, OUTPUT_CSV)
    print(f"      CSV файл: {OUTPUT_CSV}")
    print(f"\nРезультаты сохранены:")
    print(f"  • {OUTPUT_TXT}")
    print(f"  • {OUTPUT_CSV}")

if __name__ == "__main__":
    main()
